from flask import Flask, request, jsonify, send_file
from bs4 import BeautifulSoup
from docx import Document

app = Flask(__name__)

# Function to convert HTML to Word document
def convert_html_to_word(html_content, output_file="output.docx"):
    try:
        soup = BeautifulSoup(html_content, "html.parser")
        document = Document()

        # Extract and format HTML elements for Word
        for tag in soup.find_all(["h1", "h2", "h3", "p", "ul", "ol", "li"]):
            if tag.name in ["h1", "h2", "h3"]:
                document.add_heading(tag.text, level=int(tag.name[1]))
            elif tag.name == "p":
                document.add_paragraph(tag.text)
            elif tag.name in ["ul", "ol"]:
                for li in tag.find_all("li"):
                    document.add_paragraph(li.text, style="List Bullet" if tag.name == "ul" else "List Number")

        # Save the Word document
        document.save(output_file)
        return output_file
    except Exception as e:
        print(f"Error converting HTML to Word: {e}")
        return None

# Default route for root endpoint
@app.route("/")
def home():
    return "Welcome to the HTML to Word Converter API! Use the /convert endpoint to upload HTML and generate a Word document."

# API Endpoint to convert HTML to Word
@app.route("/convert", methods=["POST"])
def convert_to_doc():
    # Get HTML content from the POST request
    html_content = request.get_json().get("html_content")
    if not html_content:
        return jsonify({"error": "No HTML content provided"}), 400

    # Convert HTML to Word
    output_file = "generated_document.docx"
    result = convert_html_to_word(html_content, output_file)
    if result:
        return send_file(output_file, as_attachment=True)
    else:
        return jsonify({"error": "Failed to generate Word document"}), 500

if __name__ == "__main__":
    app.run(host="0.0.0.0", port=5000)
